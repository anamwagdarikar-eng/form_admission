import csv
import io
import os
from datetime import date

import qrcode
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font

st.set_page_config(page_title="FY Engineering Admission Form", page_icon="📝", layout="wide", initial_sidebar_state="collapsed")
RECORDS_FILE = "admission_records.csv"
PAYMENT_QR_DATA = "upi://pay?pa=vvpiet@upi&pn=VVP%20Institute%20of%20Engineering%20and%20Technology&cu=INR"

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:wght@400;700&family=Source+Sans+3:wght@400;600;700&display=swap');
:root { --ink:#1c2429; --paper:#fffdf7; --accent:#b52f3d; }
.stApp { background:#e8e2d9; color:var(--ink); font-family:'Source Sans 3',sans-serif; }
.block-container { max-width:1120px; padding:2rem 1rem 4rem; }
.paper { background:var(--paper); border:1px solid #bbb5aa; box-shadow:0 12px 34px #37302a24; padding:2rem 2.5rem; }
.masthead { display:grid; grid-template-columns:92px 1fr 92px; gap:18px; align-items:center; border-bottom:2px solid var(--ink); padding-bottom:10px; text-align:center; }
.seal { width:78px; height:78px; border:3px double var(--ink); border-radius:50%; display:grid; place-items:center; font:700 11px 'Libre Baskerville'; line-height:1.1; }
.masthead h1 { font:700 18px 'Libre Baskerville'; margin:0 0 5px; }.masthead p { margin:2px 0; font-size:11px; line-height:1.2; }
.photo { border:1px solid var(--ink); height:96px; display:grid; place-items:center; font-size:11px; }
.title { text-align:center; border-bottom:1px solid var(--ink); padding:8px 0 6px; margin-bottom:8px; font:700 14px 'Libre Baskerville'; }
.mr { color:#7d2630; font-size:12px; font-weight:600; display:block; line-height:1.1; }.paper hr { border:0; border-top:1px solid #777; margin:9px 0; }
.paper [data-testid="stForm"] { border:0; padding:0; }.paper label { font-size:12px!important; color:var(--ink)!important; }
.paper .stTextInput, .paper .stNumberInput, .paper .stDateInput, .paper .stSelectbox, .paper .stTextArea { margin-bottom:-5px; }
.paper input, .paper textarea, .paper div[data-baseweb="select"] { border-radius:0!important; background:#fff!important; border-color:#879097!important; }
.paper .stButton button { border-radius:2px; background:var(--accent); color:white; border:0; font-weight:700; }
@media (max-width:700px) { .paper { padding:1rem .85rem; }.masthead { grid-template-columns:58px 1fr 58px; gap:7px; }.seal { width:54px; height:54px; font-size:8px; }.masthead h1 { font-size:12px; }.masthead p { font-size:8px; }.photo { height:70px; font-size:9px; } }
</style>
""", unsafe_allow_html=True)


def label(text, marathi):
    return f"{text}  •  {marathi}"


def text_field(title, marathi, key, **kwargs):
    return st.text_input(label(title, marathi), key=key, **kwargs)


def marks_table(prefix, title):
    st.markdown(f"**{title}**")
    columns = st.columns([1.2, 1, 1.8, 1, 1, .7, 1, .7])
    headers = ["", "Physics", "Chemistry / ______", "Maths", "PCM Total", "%", "Grand Total", "%"]
    for column, header in zip(columns, headers):
        column.markdown(f"<div style='font-size:11px;font-weight:700;text-align:center;height:28px'>{header}</div>", unsafe_allow_html=True)
    values = {}
    for row in ["Marks Obtained", "Out of"]:
        row_columns = st.columns([1.2, 1, 1.8, 1, 1, .7, 1, .7])
        row_columns[0].markdown(f"<div style='font-size:11px;padding-top:8px'>{row}</div>", unsafe_allow_html=True)
        for index, column in enumerate(row_columns[1:]):
            values[f"{row}_{index}"] = column.text_input(f"{prefix}_{row}_{index}", label_visibility="collapsed", key=f"{prefix}_{row}_{index}")
    return values


def load_records():
    if not os.path.exists(RECORDS_FILE):
        return []
    with open(RECORDS_FILE, newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def append_record(record):
    records = load_records()
    fields = list(record)
    with open(RECORDS_FILE, "w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fields)
        writer.writeheader()
        writer.writerows(records + [record])


def word_download(record):
    document = Document()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("VIDYA VIKAS PRATISHTHAN'S\nVVP INSTITUTE OF ENGINEERING & TECHNOLOGY, SOLAPUR")
    run.bold = True
    document.add_paragraph("APPLICATION FORM FOR ADMISSION TO FIRST YEAR ENGINEERING COURSE").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Student details")
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for key, value in record.items():
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)
    document.add_page_break()
    document.add_heading("For Office Use Only", level=1)
    office = document.add_table(rows=1, cols=3)
    for cell, text in zip(office.rows[0].cells, ["Designation", "Remarks / Particulars", "Signature and Date"]):
        cell.text = text
    for role in ["Admission Committee", "Account", "Registrar", "Principal"]:
        office.add_row().cells[0].text = role
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def excel_download(records):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Admissions"
    if records:
        sheet.append(list(records[0]))
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for record in records:
            sheet.append([record.get(key, "") for key in records[0]])
    else:
        sheet.append(["No admissions submitted yet"])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def payment_qr_bytes():
    output = io.BytesIO()
    qrcode.make(PAYMENT_QR_DATA).save(output, format="PNG")
    return output.getvalue()


with st.sidebar:
    st.header("Reference")
    st.image("admission_form_FY.jpeg", width="stretch")
    st.caption("Original paper form supplied for layout reference.")

st.markdown('<div class="paper">', unsafe_allow_html=True)
st.markdown("""
<div class="masthead"><div class="seal">VVP<br>INSTITUTE<br>SEAL</div><div><p>VIDYA VIKAS PRATISHTHAN'S</p><h1>VVP INSTITUTE OF ENGINEERING &amp; TECHNOLOGY, SOLAPUR</h1><p>NAAC Accredited &amp; ISO 9001 : 2015 Certified Institute</p><p>Approved by AICTE New Delhi &amp; Affiliated to DBATU, Lonere</p><p><b>72 / 2 B, Pratapnagar, Soregaon-Dongare Road, Solapur - 413008</b></p><p>Phone: 8380305555 &nbsp; Email: vvpiet@rediffmail.com &nbsp; Website: www.vvpengg.org</p></div><div class="photo">PHOTO</div></div>
<div class="title">APPLICATION FORM FOR ADMISSION TO FIRST YEAR ENGINEERING COURSE (20 - 20 &nbsp;&nbsp; )</div>
""", unsafe_allow_html=True)
st.caption("Complete the form, choose a payment method, then submit to enable downloads.")

with st.form("admission_form"):
    st.markdown("**1. Full Name: (In Block Letters)** <span class='mr'>पूर्ण नाव (ठळक अक्षरात)</span>", unsafe_allow_html=True)
    first, father, mother = st.columns(3)
    full_name = first.text_input(label("Candidate full name", "उमेदवाराचे पूर्ण नाव"), key="full_name")
    father_name = father.text_input(label("Father's full name", "वडिलांचे पूर्ण नाव"), key="father_name")
    mother_name = mother.text_input(label("Mother's Name", "आईचे नाव"), key="mother_name")
    st.markdown("**2. Address for Correspondence** <span class='mr'>पत्रव्यवहाराचा पत्ता</span>", unsafe_allow_html=True)
    address = st.text_area(label("Address", "पत्ता"), key="address", height=68)
    pin_code = text_field("Pin Code", "पिन कोड", "pin_code")
    st.markdown("**3. Permanent Address** <span class='mr'>कायमचा पत्ता</span>", unsafe_allow_html=True)
    permanent_address = st.text_area(label("Permanent address", "कायमचा पत्ता"), key="permanent_address", height=68)
    permanent_pin = text_field("Permanent address Pin Code", "कायमचा पिन कोड", "permanent_pin")
    contact = st.columns(4)
    residence_phone = contact[0].text_input(label("Res. Phone No.", "निवासी दूरध्वनी क्र."), key="residence_phone")
    parent_phone = contact[1].text_input(label("Parent No.", "पालकांचा क्र."), key="parent_phone")
    candidate_no = contact[2].text_input(label("Candidate No.", "उमेदवार क्र."), key="candidate_no")
    email = contact[3].text_input(label("E-Mail ID", "ई-मेल आयडी"), key="email")
    identity = st.columns(3)
    sex = identity[0].selectbox(label("Sex", "लिंग"), ["Select", "Male", "Female"], key="sex")
    nationality = identity[1].text_input(label("Nationality", "राष्ट्रीयत्व"), key="nationality")
    religion = identity[2].text_input(label("Religion", "धर्म"), key="religion")
    caste = text_field("Caste: OPEN / SC / ST / SBC / VJNT / OBC / SEBC", "जात", "caste")
    sub_caste = text_field("11. (a) Sub Caste", "पोटजात", "sub_caste")
    st.markdown("**12. Details of qualifying examination (HSC / Diploma / B.Sc.) Passed. Year of Passing** <span class='mr'>पात्रता परीक्षेचा तपशील व उत्तीर्ण वर्ष</span>", unsafe_allow_html=True)
    passing_year = text_field("Year of Passing", "उत्तीर्ण वर्ष", "passing_year")
    qualifying_marks = marks_table("qualifying", "Qualifying examination marks")
    st.markdown("**13. Details of qualifying examination MHT-CET-20 / JEE (Main) - 20** <span class='mr'>MHT-CET / JEE (मुख्य) परीक्षेचा तपशील</span>", unsafe_allow_html=True)
    entrance_marks = marks_table("entrance", "Entrance examination marks")
    st.markdown("**14. Details of SSC Examination Passed:** <span class='mr'>SSC परीक्षेचा तपशील</span>", unsafe_allow_html=True)
    ssc = st.columns([2, 2, 1.4, 1.2])
    ssc_school = ssc[0].text_input(label("Name of School", "शाळेचे नाव"), key="ssc_school")
    ssc_board = ssc[1].text_input(label("Name of Board", "बोर्डचे नाव"), key="ssc_board")
    ssc_year = ssc[2].text_input(label("Year of Passing", "उत्तीर्ण वर्ष"), key="ssc_year")
    ssc_marks = ssc[3].text_input(label("% of Marks", "गुणांची टक्केवारी"), key="ssc_marks")
    st.markdown("**15. Course / Branch Selected for F. Y. B. Tech.** <span class='mr'>प्रथम वर्ष बी.टेक. साठी निवडलेला अभ्यासक्रम / शाखा</span>", unsafe_allow_html=True)
    branches = ["Artificial Intelligence & Data Science", "Computer Science & Engineering", "Civil Engineering", "Electrical Engineering", "Electronic & Telecommunication Engineering", "Mechanical Engineering"]
    branch_cols = st.columns(2)
    selected_branches = []
    for index, branch in enumerate(branches):
        if branch_cols[index % 2].checkbox(f"{index + 1}. {branch}", key=f"branch_{index}"):
            selected_branches.append(branch)
    st.markdown("**Payment Details** <span class='mr'>देयक तपशील</span>", unsafe_allow_html=True)
    payment_mode = st.radio(label("Payment mode", "देयक पद्धत"), ["Online", "Offline"], horizontal=True, key="payment_mode")
    if payment_mode == "Online":
        st.image(payment_qr_bytes(), width=180, caption="Scan to pay admission fees")
        payment_reference = text_field("Online transaction reference", "ऑनलाईन व्यवहार क्रमांक", "payment_reference")
    else:
        payment_reference = text_field("Offline receipt number", "ऑफलाईन पावती क्रमांक", "payment_reference")
    submitted = st.form_submit_button("Submit Admission Form")

if submitted:
    if not full_name.strip():
        st.error("Candidate full name is required.")
    elif payment_mode == "Offline" and not payment_reference.strip():
        st.error("Enter the offline receipt number before submitting.")
    else:
        record = {"Submitted On": str(date.today()), "Full Name": full_name, "Father Name": father_name, "Mother Name": mother_name, "Address": address, "Pin Code": pin_code, "Permanent Address": permanent_address, "Permanent Pin": permanent_pin, "Residence Phone": residence_phone, "Parent Phone": parent_phone, "Candidate No": candidate_no, "Email": email, "Sex": sex, "Nationality": nationality, "Religion": religion, "Caste": caste, "Sub Caste": sub_caste, "Passing Year": passing_year, "Qualifying Marks": str(qualifying_marks), "Entrance Marks": str(entrance_marks), "SSC School": ssc_school, "SSC Board": ssc_board, "SSC Year": ssc_year, "SSC Marks": ssc_marks, "Selected Branches": ", ".join(selected_branches), "Payment Mode": payment_mode, "Payment Reference": payment_reference}
        append_record(record)
        st.session_state["last_record"] = record
        st.success("Admission submitted and added to the admission register.")

st.markdown("</div>", unsafe_allow_html=True)
last_record = st.session_state.get("last_record")
if last_record:
    st.download_button("Download Admission Form (Word)", word_download(last_record), "admission_form.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.subheader("Admission Register")
records = load_records()
if records:
    st.dataframe(records, use_container_width=True, hide_index=True)
    st.download_button("Download Admission Register (Excel)", excel_download(records), "admission_register.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
else:
    st.info("No admission records have been submitted yet.")